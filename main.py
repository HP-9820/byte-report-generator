import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))


import time
import json
import random
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.data_loader import load_and_process_data
from datetime import datetime
import re

# Import for spider chart
import numpy as np
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
from io import BytesIO


# --- CONFIGURATION ---
DEEPSEEK_API_KEY = "sk-5fc2fb744dac49fd821977bdfea028b2"
INPUT_FOLDER = "input"
OUTPUT_FOLDER = "output"
EXPLAINABILITY_FOLDER = "explainability_logs"

# --- MODEL: DEEPSEEK CHAT ---
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

# =========================
# INTEREST LEVEL LOGIC
# =========================
INTEREST_LEVEL_SYNONYMS = {
    "HIGH": [
        "Highly Developed Interest",
        "Strong Mastery and Enthusiasm",
        "Advanced Engagement and Confidence",
        "Consistent Strength and Passion"
    ],
    "STRONG": [
        "Showing Strong Interest",
        "Positive and Active Engagement",
        "Well-Developed Interest",
        "Growing Strength in This Area"
    ],
    "DEVELOPING": [
        "Developing Engagement",
        "Emerging Interest",
        "Building Confidence Gradually",
        "Progressing With Guidance"
    ],
    "EXPLORING": [
        "Exploring Preferences",
        "Early Curiosity Observed",
        "Initial Engagement Noted",
        "Beginning to Show Interest"
    ],
    "EARLY": [
        "Early Exposure",
        "Limited Experience So Far",
        "Initial Familiarity Developing",
        "Requires More Opportunities"
    ],
    "NA": ["Not Assessed Yet"]
}

def get_interest_label(score):
    if score is None:
        return random.choice(INTEREST_LEVEL_SYNONYMS["NA"])
    if score >= 85:
        return random.choice(INTEREST_LEVEL_SYNONYMS["HIGH"])
    elif score >= 60:
        return random.choice(INTEREST_LEVEL_SYNONYMS["STRONG"])
    elif score >= 40:
        return random.choice(INTEREST_LEVEL_SYNONYMS["DEVELOPING"])
    elif score >= 20:
        return random.choice(INTEREST_LEVEL_SYNONYMS["EXPLORING"])
    else:
        return random.choice(INTEREST_LEVEL_SYNONYMS["EARLY"])


# =========================
# DOCX TABLE STYLING HELPERS
# =========================
def style_table_header(row, bg_color="1F5E78", text_color=RGBColor(255, 255, 255)):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg_color)
        tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = text_color

def set_column_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

def add_formatted_text_with_bold(cell, text):
    """
    Adds text to a cell with **bold** markdown-style formatting.
    Converts **text** into actual bold runs.
    """
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(9)
        if i % 2 == 1:
            run.bold = True

def format_all_points(cell, text):
    """
    Displays support activities as normal paragraph text.
    Removes numbering/bullets formatting.
    """
    cell._element.clear_content()

    cleaned_text = str(text).strip()

    # remove numbering like 1. 2. 3.
    cleaned_text = re.sub(r'\d+\.\s*', '', cleaned_text)

    # replace semicolons with proper sentence spacing
    cleaned_text = cleaned_text.replace(';', '. ')

    p = cell.add_paragraph()
    run = p.add_run(cleaned_text)
    run.font.size = Pt(8.5)

def format_future_possibilities(cell, future_possibilities_text):
    """
    Renders the future_possibilities field as 3 labelled tracks:
    Government, Private, Entrepreneurship — each on its own line.
    Input format expected:
    Government: "X, Y, Z" | Private: "A, B, C" | Entrepreneurship: "D, E, F"
    """
    cell._element.clear_content()

    if not future_possibilities_text or str(future_possibilities_text).strip().lower() in ["null", "none", ""]:
        p = cell.add_paragraph()
        run = p.add_run("Not identified")
        run.font.size = Pt(8.5)
        run.italic = True
        return

    # Split by pipe separator
    tracks = [t.strip() for t in str(future_possibilities_text).split('|') if t.strip()]

    for track in tracks:
        if ':' in track:
            label, careers = track.split(':', 1)
            label = label.strip()
            careers = careers.strip().strip('"').strip("'").strip()
        else:
            label = ""
            careers = track.strip()

        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(1)

        if label:
            label_run = p.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(8.5)
            label_run.font.color.rgb = RGBColor(31, 94, 120)

        careers_run = p.add_run(careers)
        careers_run.font.size = Pt(8.5)

def create_bordered_paragraph(cell, text, bold=False):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if bold:
        run.bold = True
    tc_pr = cell._tc.get_or_add_tcPr()
    for border_side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_pr.append(border)


# =========================
# SPIDER CHART GENERATION
# =========================
def create_spider_chart(student_data, class_avg_data):
    try:
        activities = []
        student_scores = []
        class_scores = []

        for activity, details in student_data["activities"].items():
            score = details.get("average")
            if activity.lower() == "others" and (score is None or score == 0):
                continue
            activities.append(activity)
            student_scores.append(round(score) if score is not None else 0)
            class_item = class_avg_data.get(activity, 0)
            if isinstance(class_item, dict):
                class_score = class_item.get("average", 0)
            else:
                class_score = class_item
            class_scores.append(round(class_score) if class_score else 0)

        if len(activities) < 2:
            print("   ⚠️ Not enough activities for spider chart (need at least 2)")
            return None

        num_vars = len(activities)
        angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
        student_scores_plot = student_scores + student_scores[:1]
        class_scores_plot = class_scores + class_scores[:1]
        angles_plot = angles + angles[:1]

        fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
        fig.patch.set_alpha(0)
        ax.set_facecolor('none')
        ax.set_theta_offset(np.pi / 2)
        ax.set_theta_direction(-1)
        ax.set_xticks(angles)
        ax.set_xticklabels(activities, size=9, weight='bold')
        ax.tick_params(axis='x', pad=25)
        ax.set_rlabel_position(0)
        ax.set_yticks([25, 50, 75, 100])
        ax.set_yticklabels(['25', '50', '75', '100'], size=8, color='gray')
        ax.set_ylim(0, 100)
        ax.grid(True, linestyle='--', color='gray', alpha=0.5, linewidth=1.2)
        ax.spines['polar'].set_linewidth(2.0)
        ax.spines['polar'].set_color('#333333')
        ax.plot(angles_plot, student_scores_plot, linewidth=2.5, marker='o',
                label='Student average', color='#4472C4', markersize=6)
        ax.fill(angles_plot, student_scores_plot, alpha=0.25, color='#4472C4')
        ax.plot(angles_plot, class_scores_plot, linewidth=2.5, marker='o',
                label='Class average', color='#ED7D31', markersize=6)
        ax.fill(angles_plot, class_scores_plot, alpha=0.15, color='#ED7D31')
        ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.1), frameon=False, fontsize=9)

        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=200, bbox_inches='tight',
                    transparent=True, pad_inches=0.1)
        img_buffer.seek(0)
        return img_buffer

    except Exception as e:
        print(f"   ⚠️ Error creating spider chart: {e}")
        return None
    finally:
        plt.close('all')


# =========================
# LLM ANALYSIS
# =========================
def generate_llm_analysis(student_data, student_name, class_name):
    # 1. Prepare Data with Calculated Synonyms
    raw_student_data = "Activity | Score | Interest Level | Feedback\n"
    for activity, details in student_data["activities"].items():
        score = round(details["average"]) if details.get("average") is not None else None
        interest_label = get_interest_label(score)
        score_text = f"{score}%" if score is not None else "-"
        feedbacks = "; ".join(details["feedbacks"][:5]) if details.get("feedbacks") else "No specific feedback recorded"
        raw_student_data += f"{activity} | {score_text} | {interest_label} | {feedbacks}\n"

    # 2. Tone Block Logic
    tone_block = "BLOCK P – Primary Classes 3–4"
    try:
        class_num = int(re.search(r"\d+", str(class_name)).group())
        if class_num in [3, 4]:
            tone_block = "BLOCK P – PRIMARY (Classes 3–4)"
        elif class_num in [5, 6]:
            tone_block = "BLOCK U – UPPER PRIMARY (Classes 5–6)"
        elif class_num in [7, 8, 9]:
            tone_block = "BLOCK M – MIDDLE SCHOOL (Classes 7–9)"
    except:
        pass

    prompt = f"""
<SYSTEM PROMPT – PARENT CO-CURRICULAR REPORT>

You are an expert Educational Consultant designing parent-friendly,
practical, and class-appropriate co-curricular career recommendations
for Indian families.

Your goal is to help parents clearly understand:
• What the child is naturally inclined towards
• What practical steps parents can take next
• What future directions may gradually open up (without pressure)

LANGUAGE & STYLE RULES (STRICT)
• Use only simple, everyday words
• Write warmly and encouragingly
• Never use the student's name
• Never use pronouns (he / she / they)
• Always write in present tense
• Write for Indian parents (realistic, grounded)

CONTEXT
Class Level: {class_name}
Tone & Career Block: {tone_block}

RAW STUDENT DATA (pipe-delimited):
{raw_student_data}

-------------------------------------------------
CRITICAL RULES
-------------------------------------------------

1) INTEREST LEVELS (NON-NEGOTIABLE)
• Interest levels are already provided
• Use the exact wording from the data
• Do NOT rephrase, explain, or substitute

-------------------------------------------------
2) OBSERVATIONS (MANDATORY)
• Every subject and activity in the data must appear
• Write 1–2 simple sentences per area
• Focus on effort, engagement, curiosity, and learning behavior
• Do NOT compare with other students

-------------------------------------------------
3) RECOMMENDATIONS (TOP 2 ONLY – WITH PARENT ACTIONABILITY)

Select ONLY the top 2 highest engagement areas.

-------------------------------------------------
3.1 SELECTION RULE
-------------------------------------------------
• Recommendations MUST be based on the TOP 2 areas by engagement score — this includes Academics if it ranks in the top 2.
• If Academics has the highest engagement score, it MUST appear as the first recommendation.
• For Academics recommendations, focus on the specific subjects where the student excels (e.g., Mathematics, Science) and career paths those subjects lead to.
• Do NOT skip or deprioritize Academics simply because it is an academic subject.
• Keep each recommendation header focused on the single area name only (e.g., "Academics", "Computers").

-------------------------------------------------
3.2 CLASS-BASED CAREER GUIDANCE
-------------------------------------------------

For Classes 3–5 (ALL areas including Academics):
• Do NOT mention specific job titles
• Focus on exposure, confidence, habits, and skill discovery
• Career direction must remain broad and future-facing
• For Academics: Highlight the joy of learning, curiosity, and subject strengths
• Mention broad fields like "science and discovery", "numbers and logic",
  "reading and writing" — NOT specific roles like "Doctor" or "Engineer"
• Encourage subject-based clubs, olympiads, or hobby reading

For Classes 6–7 (ALL areas including Academics):
• Introduce career fields (not final roles)
• Explain how current academic strengths connect to career fields
• For Academics (Science): Mention fields like medicine, engineering, research
• For Academics (Maths): Mention fields like finance, technology, data
• For Academics (Social): Mention fields like law, civil services, journalism
• For Academics (Languages): Mention fields like media, writing, communications
• Suggest structured learning, subject olympiads, clubs, or guided practice

For Classes 8–9 (ALL areas including Academics):
• Mention specific career paths where relevant
• For Academics: Clearly connect strong subjects to specific career streams
  - Science → Medicine, Engineering, Research, Biotechnology
  - Maths → Data Science, Finance, Software, Architecture
  - Social → Civil Services, Law, Journalism, International Relations
  - Languages → Media, Publishing, Diplomacy, Content Creation
• Keep tone supportive and non-pressurizing
• Emphasize stream selection readiness (Science/Commerce/Arts) where appropriate
• Reinforce that strong academics opens doors across ALL career fields

-------------------------------------------------
3.3 FORMATTING REQUIREMENTS (STRICT)
-------------------------------------------------

CAREER PATHWAYS FORMAT:
• Start with a brief intro sentence (1 line)
• List 3-5 specific career paths in BOLD using **Career Name** format
• Keep descriptions practical and realistic
• Connect careers logically to the student's interests
• ALWAYS pick careers from the CAREER KNOWLEDGE BASE below
• Never invent career names not listed in the knowledge base

CAREER SELECTION BY CLASS LEVEL (STRICT):

For Classes 3–5:
• Pick only BROAD, easy-to-understand career categories
• Avoid technical job titles like "DevOps Engineer" or "Actuary"
• Use approachable titles: **Sports Coach**, **Doctor**, **Artist**, **Teacher**
• Focus on what the child DOES, not the job title

For Classes 6–7:
• Introduce career fields with simple names
• Acceptable: **Software Developer**, **Sports Manager**, **Graphic Designer**
• Avoid highly technical titles like **Quant Analyst**, **IES Officer**
• Briefly explain what each career involves in one phrase

For Classes 8–9:
• Use specific career titles from the knowledge base
• Acceptable: **AI/ML Engineer**, **Civil Servant (IAS/IPS/IFS)**, **Chartered Accountant**
• Connect each career to a specific skill or subject strength
• Mention entrance exams or degree paths where relevant (JEE, NEET, CLAT, UPSC)

CAREER CATEGORY GUIDE (use to pick age-appropriate careers):

EASY (Classes 3–5): Sports Coach, Doctor, Artist, Teacher, Chef, Farmer, Pilot,
Nurse, Designer, Fitness Trainer

MODERATE (Classes 6–7): Software Developer, Graphic Designer, Sports Manager,
Game Designer, Content Creator, Marketing Manager, School Teacher,
Environmental Scientist, Fashion Designer, Event Planner, Animator/VFX Artist

ADVANCED (Classes 8–9): AI/ML Engineer, Civil Servant (IAS/IPS/IFS),
Chartered Accountant, Investment Banker, Robotics/Mechatronics Engineer,
Lawyer/Advocate, Cybersecurity Analyst, FinTech Product Manager,
Clinical Psychologist, Actuary, Public Policy Analyst, Commercial Pilot,
Indian Forest Service Officer, Biotechnologist, Healthcare Administrator

Example Format (Class 8–9):
"Strong academic performance opens many exciting career paths. Exploring
**AI/ML Engineer**, **Biotechnologist**, **Civil Servant (IAS/IPS/IFS)**,
or **Chartered Accountant** is very possible. Each of these paths rewards
hard work, curiosity, and strong subject knowledge."

Example Format (Class 6–7):
"This growing interest in computers can lead to exciting fields.
Exploring **Software Developer**, **Game Designer**, **Graphic Designer**,
or **Content Creator** could be great directions. These fields reward
creativity and logical thinking."

Example Format (Class 3–5):
"This love for drawing can open many wonderful paths.
Exploring **Artist**, **Designer**, **Teacher**, or **Chef** is possible.
Trying new activities and building confidence is what matters most now."

-------------------------------------------------
PARENT ACTIONABILITY RULE (MANDATORY):
-------------------------------------------------

For EVERY recommendation, provide EXACTLY 3 parent actions in numbered format.

Each action MUST:
• Be specific and practical (not generic motivation)
• Appropriate for the class level
• Directly connected to the student's interest and engagement data
• Easy for Indian parents to act on within the next 6–12 months
• Mention the type of class, activity, or routine
• Clearly explain what skill the action strengthens
• Avoid vague phrases like "encourage interest", "explore more", "provide exposure"

NUMBERED FORMAT (MANDATORY):
Write actions as: "1. [Action]; 2. [Action]; 3. [Action]"

Example Format:
"1. Enroll in a structured badminton academy or shuttle coaching program to strengthen specific skills; 2. Encourage participation in school or inter-school sports competitions to build competitive spirit and teamwork; 3. Provide opportunities to learn about sports nutrition and fitness routines to support physical development."

CLASS-WISE EXPECTATIONS:

For Classes 3–5:
• Focus on habit building, confidence, and exposure
• Suggest hobby classes, simple home routines, and school-level participation
• For Academics: Suggest subject reading books, math puzzles, science
  experiment kits, story writing habits
• Avoid exams, certifications, or career pressure

For Classes 6–7:
• Focus on structured learning and skill strengthening
• Suggest guided classes, small projects, competitions, or clubs
• For Academics: Suggest subject olympiads (Math/Science), debate clubs,
  science fairs, book clubs, quiz competitions
• Introduce career fields indirectly through subject strengths

For Classes 8–9:
• Focus on preparation and clarity
• For Academics: Suggest foundation courses (JEE/NEET/CLAT foundation),
  subject mentoring, competitive exam exposure, stream selection guidance
• Suggest portfolios, mentoring, or academies where relevant
• Clearly connect strong academic subjects to future career paths and
  stream choices (Science / Commerce / Arts)

-------------------------------------------------
5) HOW PARENTS CAN HELP: PARENT TIPS
-------------------------------------------------
Provide exactly 3 high-quality tips in the "parent_tips" field.
Use this specific style (combining activity context with practical timing):

Style Examples for Academics as top interest:
1. Weekend Exploration (Academics - Science): Dedicate a weekend afternoon
   to a fun science experiment at home using simple kitchen materials to
   strengthen curiosity and hands-on learning in science.
2. Short Daily Practice (Computers): Encourage 15 minutes of typing practice
   or a simple coding activity to build logical thinking alongside academic strengths.
3. Encourage and Talk: Dedicate 10–15 minutes weekly to ask about favorite
   subjects, what new things were learned, and celebrate small academic wins.

Style Examples for non-academic top interest:
1. Weekend Exploration (Computers): Dedicate a weekend afternoon to a fun,
   structured online lesson that teaches basic data organization or typing
   skills, connecting it to the interest in data entry.
2. Short Daily Practice (Drawing): Encourage 15 minutes of sketching simple
   objects or geometric shapes in the evening to maintain the interest in
   exact measurements and shading techniques.
3. Encourage and Talk: Dedicate 10–15 minutes weekly to talk about what
   excites your child and offer your appreciation for their efforts.

Requirements for the 3 tips:
• Tip 1: "Weekend Exploration" — related to the child's TOP interest area.
  For Academics: connect to the strongest subject (Science/Maths/Social/Languages).
  For others: connect to the specific activity.
• Tip 2: "Short Daily Practice" (15 mins) — related to the SECOND interest area.
• Tip 3: "Encourage and Talk" — 10–15 minutes weekly, warm and encouraging tone.

CLASS-WISE TIP CALIBRATION:
• Classes 3–5: Simple, fun, home-based activities. No exam pressure.
• Classes 6–7: Structured practice, clubs, olympiads, online learning.
• Classes 8–9: Foundation courses, competitive exam prep, mentoring sessions.

-------------------------------------------------
6) CONCLUSION
-------------------------------------------------
• Write a warm, reassuring 3–4 sentence summary
• Emphasize steady growth and natural strengths
• Reinforce that exploration is healthy at this stage
• Mention both top interest areas briefly
• Do NOT mention the class name directly
• End with an encouraging note for parents


CAREER SELECTION LOGIC (STRICT — FOLLOW EXACTLY):

Step 1: Look at ALL activities and their engagement scores (including Academics).
Step 2: Identify the SINGLE activity or subject with the HIGHEST engagement score.
        — If Academics has the highest score → use the specific Academic subject
          that is strongest (Science / Mathematics / Social Studies / Languages).
        — If a non-academic activity has the highest score → use that activity.
Step 3: Select careers ONLY from that one winning category below.
        Do NOT mix careers from multiple categories.
        Do NOT create combinations.
Step 4: Always output exactly 3 tracks: Government, Private, Entrepreneurship.
Step 5: List 3–4 careers per track as plain comma-separated text.
Step 6: If no dominant activity is identified → return null.

CAREER PATHS BY CATEGORY:

SPORTS:
Sports Scientist
Physiotherapist
Sports Nutritionist
Strength & Conditioning Coach
Exercise Physiologist
Sports Data Analyst
Performance Analyst
Sports Statistics Expert
Sports Journalist
Sports Commentator
Sports Content Creator
Sports PR Manager
Sports Administrator
Sports Policy Analyst
Sports Management Professional
Community Sports Officer

LIBRARY:
Author
Writer
Editor
Publisher
Journalist
Scriptwriter
Content Strategist
Historian
Civil Services Officer
Political Analyst
Policy Researcher
Think Tank Researcher
Research Scholar
Academic Scientist
Science Communicator
Science Writer
Economist
Actuary
Data Research Analyst
Financial Analyst

COMPUTERS:
Software Engineer
Data Scientist
AI/ML Engineer
Game Developer
Quant Analyst
Robotics Engineer
Bioinformatics Specialist
Environmental Tech Analyst
Health Tech Developer
Technical Writer
UX Writer
Product Manager
Digital Marketer
Instructional Designer
Civic Tech Specialist
GovTech Consultant
Digital Policy Analyst
Cyber Law Professional

ARTS:
Filmmaker
Lyricist
Screenwriter
Theatre Artist
Creative Director
Fashion Designer
Cultural Researcher
Interior Designer
Heritage Conservationist
Medical Illustrator
Industrial Designer
Product Designer
Architect
UI/UX Designer
Animation & VFX Artist
Game Designer

MUSIC:
Lyricist
Music Journalist
Radio Jockey
Podcast Host
Music Educator
Sound Engineer
Music Producer
Audio Technology Specialist

DANCE:
Dance Therapist
Movement Analyst
Fitness Choreographer
Cultural Program Coordinator
Dance Historian
Arts Administrator

RULES:
• Output must be plain text — no bullets, no formatting, no JSON inside this field
• Career names must be simple and recognizable to Indian parents
• Do NOT invent careers outside the lists above
• Do NOT blend careers from multiple categories
• The winning category is determined ONLY by the highest engagement score


-------------------------------------------------
OUTPUT FORMAT (JSON ONLY)
-------------------------------------------------
{{
  "observations": [
    {{
      "activity": "Name",
      "engagement": 90,
      "interest_level": "Exact String From Data",
      "observation": "Text"
    }}
  ],
  "recommendations": [
    {{
      "area": "Name",
      "engagement": 90,
      "future_pathways": "Brief intro sentence. Mention **Career1**, **Career2**, **Career3** in bold. Additional context about skills.",
      "support_activities": "Text"
    }}
  ],
  "parent_tips": ["Tip 1", "Tip 2", "Tip 3"],
  "conclusion": "Text"
}}
"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "You are an expert Educational Consultant. Always respond in valid JSON format."},
                    {"role": "user", "content": prompt}
                ],
                response_format={'type': 'json_object'}
            )
            text = response.choices[0].message.content.strip()
            parsed = json.loads(text)
            return parsed, {"input": raw_student_data}

        except Exception as e:
            print(f"   > API Error (Attempt {attempt+1}): {e}")
            if "429" in str(e):
                wait = 30 * (attempt + 1)
                print(f"   > Rate limited. Waiting {wait}s before retry...")
                time.sleep(wait)
            else:
                time.sleep(2)

    return None, None


# =========================
# WORD REPORT CREATION
# =========================
def create_word_doc(student_name, analysis, class_name, student_data, class_avg_data,
                    output_dir=None, doc=None, save=True, is_first=False):
    if doc is None:
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
    elif not is_first:
        doc.add_page_break()

    target_folder = output_dir if output_dir else OUTPUT_FOLDER

    # ========================================
    # STUDENT INFO TABLE
    # ========================================
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    set_table_borders(info_table)
    set_column_widths(info_table, [Inches(1.2), Inches(5.8)])

    info_table.rows[0].cells[0].text = "Name"
    info_table.rows[0].cells[1].text = student_name
    style_table_header(info_table.rows[0], bg_color="1F5E78")

    info_table.rows[1].cells[0].text = "Class"
    info_table.rows[1].cells[1].text = class_name

    for row in info_table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # ========================================
    # PURPOSE
    # ========================================
    purpose_heading = doc.add_heading("Purpose:", level=2)
    purpose_heading.style.font.size = Pt(11)
    purpose_heading.style.font.bold = True
    purpose_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    purpose_heading.paragraph_format.space_before = Pt(2)
    purpose_heading.paragraph_format.space_after = Pt(1)

    purpose_text = doc.add_paragraph(
        '"Skill Assessment Program - Building Strengths Beyond Academics "\n'
        'This report shows your child\'s interests and abilities in different school activities. '
        'With your (parents) support and our guidance, these strengths can grow into useful life skills. '
        'It helps parents notice early signs of talent and areas where extra care or practice can make a big difference.'
    )
    purpose_text.style.font.size = Pt(9)
    purpose_text.paragraph_format.space_after = Pt(4)

    # ========================================
    # KEY INSIGHTS + SPIDER CHART
    # ========================================
    insights_heading = doc.add_heading("Key Insights:", level=2)
    insights_heading.style.font.size = Pt(11)
    insights_heading.style.font.bold = True
    insights_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    insights_heading.paragraph_format.space_before = Pt(2)
    insights_heading.paragraph_format.space_after = Pt(1)

    insights_text = doc.add_paragraph(
        'The graph below shows how your child is doing in different areas. It tells us what they like more and where we can '
        'help them do even better. Blue represents your child, and orange represents the class average.'
    )
    insights_text.style.font.size = Pt(9)
    insights_text.paragraph_format.space_after = Pt(2)

    chart_img = create_spider_chart(student_data, class_avg_data)
    if chart_img:
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_para.paragraph_format.space_before = Pt(0)
        chart_para.paragraph_format.space_after = Pt(4)
        run = chart_para.add_run()
        run.add_picture(chart_img, width=Inches(3.2))

    # ========================================
    # OBSERVATIONS TABLE
    # ========================================
    obs_heading = doc.add_heading("Observations:", level=2)
    obs_heading.style.font.size = Pt(11)
    obs_heading.style.font.bold = True
    obs_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    obs_heading.paragraph_format.space_before = Pt(2)
    obs_heading.paragraph_format.space_after = Pt(1)

    subtitle = doc.add_paragraph("The table gives an overview of child's involvement and performance across co-curricular areas.")
    subtitle.style.font.size = Pt(9)
    subtitle.paragraph_format.space_after = Pt(2)

    sorted_obs = sorted(
        analysis.get("observations", []),
        key=lambda x: x.get('engagement', 0) if isinstance(x.get('engagement'), (int, float)) else -1,
        reverse=True
    )

    filtered_obs = []
    for obs in sorted_obs:
        if obs.get("activity", "").strip().lower() == "others":
            eng = obs.get("engagement")
            if eng is None or eng == 0:
                continue
        filtered_obs.append(obs)

    obs_table = doc.add_table(rows=1, cols=4)
    obs_table.style = 'Table Grid'
    set_table_borders(obs_table)
    set_column_widths(obs_table, [Inches(1.5), Inches(0.7), Inches(1.6), Inches(3.2)])

    hdr = obs_table.rows[0]
    hdr.cells[0].text = "Interest Area"
    hdr.cells[1].text = "Engagement (%)"
    hdr.cells[2].text = "Interest Level"
    hdr.cells[3].text = "Key Observations"
    style_table_header(hdr, bg_color="1F5E78")

    for obs in filtered_obs:
        row = obs_table.add_row().cells
        row[0].text = obs.get("activity", "")
        eng = obs.get("engagement")
        row[1].text = f"{eng}%" if isinstance(eng, (int, float)) else "-"
        row[2].text = obs.get("interest_level", "")
        row[3].text = obs.get("observation", "")
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    # ========================================
    # RECOMMENDATIONS TABLE
    # ========================================
    rec_heading = doc.add_heading("Recommendations & Growth Opportunities:", level=2)
    rec_heading.style.font.size = Pt(11)
    rec_heading.style.font.bold = True
    rec_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    rec_heading.paragraph_format.space_before = Pt(2)
    rec_heading.paragraph_format.space_after = Pt(1)

    rec_subtitle = doc.add_paragraph("Here are some ways to turn your child's interests into long-term strengths:")
    rec_subtitle.style.font.size = Pt(9)
    rec_subtitle.paragraph_format.space_after = Pt(2)

    rec_table = doc.add_table(rows=1, cols=3)
    rec_table.style = 'Table Grid'
    set_table_borders(rec_table)
    set_column_widths(rec_table, [Inches(1.5), Inches(3.5), Inches(3.0)])

    hdr = rec_table.rows[0]
    hdr.cells[0].text = "Strong Interest Area"
    hdr.cells[1].text = "Potential Careers"
    hdr.cells[2].text = "Simple Activities to Support"
    style_table_header(hdr, bg_color="1F5E78")

    recs = analysis.get("recommendations", [])[:2]

    for rec in recs:
        row = rec_table.add_row().cells

        # Column 1: Area + score
        eng = rec.get("engagement")
        score_txt = f"\n({eng}%)" if isinstance(eng, (int, float)) else ""
        row[0].text = f"{rec.get('area', '')}{score_txt}"
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8.5)

        # Column 2: Future pathways with bold careers — NO subject_career_combinations
        future_pathways = rec.get("future_pathways", "")
        add_formatted_text_with_bold(row[1], future_pathways)

        # Column 3: All 3 support activities on separate lines
        support_activities = rec.get("support_activities", "")
        format_all_points(row[2], support_activities)

    # ========================================
    # FUTURE POSSIBILITIES TABLE
    # ========================================
    fp_heading = doc.add_heading("Future Possibilities to Explore:", level=2)
    fp_heading.style.font.size = Pt(11)
    fp_heading.style.font.bold = True
    fp_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    fp_heading.paragraph_format.space_before = Pt(6)
    fp_heading.paragraph_format.space_after = Pt(1)

    fp_subtitle = doc.add_paragraph(
        "Based on the strongest area of engagement, here are career paths worth exploring across different tracks:"
    )
    fp_subtitle.style.font.size = Pt(9)
    fp_subtitle.paragraph_format.space_after = Pt(2)

    fp_table = doc.add_table(rows=1, cols=2)
    fp_table.style = 'Table Grid'
    set_table_borders(fp_table)
    set_column_widths(fp_table, [Inches(2.0), Inches(5.0)])

    fp_hdr = fp_table.rows[0]
    fp_hdr.cells[0].text = "Track"
    fp_hdr.cells[1].text = "Career Options"
    style_table_header(fp_hdr, bg_color="1F5E78")

    future_possibilities_text = analysis.get("future_possibilities", "")

    if future_possibilities_text and str(future_possibilities_text).strip().lower() not in ["null", "none", ""]:
        tracks = [t.strip() for t in str(future_possibilities_text).split('|') if t.strip()]
        for track in tracks:
            if ':' in track:
                label, careers = track.split(':', 1)
                label = label.strip()
                careers = careers.strip().strip('"').strip("'").strip()
            else:
                label = "General"
                careers = track.strip()

            fp_row = fp_table.add_row().cells

            label_para = fp_row[0].paragraphs[0]
            label_para.clear()
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGBColor(31, 94, 120)

            careers_para = fp_row[1].paragraphs[0]
            careers_para.clear()
            careers_run = careers_para.add_run(careers)
            careers_run.font.size = Pt(9)
    else:
        fp_row = fp_table.add_row().cells
        fp_row[0].text = "—"
        fp_row[1].text = "Not enough data to identify a dominant interest area."
        for cell in fp_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.italic = True

    # ========================================
    # PARENT TIPS
    # ========================================
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    tips_heading = doc.add_heading("How Parents Can Help:", level=2)
    tips_heading.style.font.size = Pt(11)
    tips_heading.style.font.bold = True
    tips_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    tips_heading.paragraph_format.space_before = Pt(2)
    tips_heading.paragraph_format.space_after = Pt(1)

    for i, tip in enumerate(analysis.get("parent_tips", []), 1):
        tip_para = doc.add_paragraph(f"{i}. {tip}")
        tip_para.style.font.size = Pt(9)
        tip_para.paragraph_format.left_indent = Inches(0.2)
        tip_para.paragraph_format.space_after = Pt(3)

    # ========================================
    # CONCLUSION
    # ========================================
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    conclusion_heading = doc.add_heading("Conclusion:", level=2)
    conclusion_heading.style.font.size = Pt(11)
    conclusion_heading.style.font.bold = True
    conclusion_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    conclusion_heading.paragraph_format.space_before = Pt(2)
    conclusion_heading.paragraph_format.space_after = Pt(1)

    conclusion_para = doc.add_paragraph(analysis.get("conclusion", ""))
    conclusion_para.style.font.size = Pt(9)
    conclusion_para.paragraph_format.space_after = Pt(4)

    # ========================================
    # NOTES
    # ========================================
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    notes_heading = doc.add_heading("Notes:", level=2)
    notes_heading.style.font.size = Pt(11)
    notes_heading.style.font.bold = True
    notes_heading.style.font.color.rgb = RGBColor(0, 51, 102)
    notes_heading.paragraph_format.space_before = Pt(2)
    notes_heading.paragraph_format.space_after = Pt(1)

    notes_table = doc.add_table(rows=2, cols=1)
    notes_table.style = 'Table Grid'
    set_table_borders(notes_table)

    for row in notes_table.rows:
        row.height = Inches(0.4)

    parents_cell = notes_table.rows[0].cells[0]
    parents_cell._element.clear_content()
    p1 = parents_cell.add_paragraph()
    run1 = p1.add_run("Parents remarks:")
    run1.bold = True
    run1.font.size = Pt(10)

    teachers_cell = notes_table.rows[1].cells[0]
    teachers_cell._element.clear_content()
    p2 = teachers_cell.add_paragraph()
    run2 = p2.add_run("Teachers' remarks:")
    run2.bold = True
    run2.font.size = Pt(10)

    if not save:
        return doc

    safe_name = "".join([c for c in student_name if c.isalnum() or c == ' ']).strip()
    output_path = f"{target_folder}/{safe_name}_Report.docx"

    try:
        doc.save(output_path)
        print(f"✓ Saved: {safe_name}_Report.docx")
    except PermissionError:
        timestamp = int(time.time())
        print(f"⚠️  WARNING: Could not save {safe_name}_Report.docx — file may be open in Word.")
        alt_path = f"{target_folder}/{safe_name}_Report_{timestamp}.docx"
        doc.save(alt_path)
        print(f"✓ Saved as: {safe_name}_Report_{timestamp}.docx instead")

    return doc


# =========================
# MAIN
# =========================
def main():
    print("--- Starting Report Generation ---")

    # TEST MODE CONFIGURATION
    TEST_MODE = True   # Set to False to process all students
    MAX_STUDENTS = 5   # Number of students to process in test mode

    if not os.path.exists(INPUT_FOLDER):
        print(f"❌ Error: {INPUT_FOLDER} folder not found!")
        return

    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        print(f"✓ Created {OUTPUT_FOLDER} folder")

    files = [
        f for f in os.listdir(INPUT_FOLDER)
        if f.endswith(('.xlsx', '.xls', '.csv')) and not f.startswith('~')
    ]

    if not files:
        print(f"❌ Error: No Excel/CSV files found in {INPUT_FOLDER}")
        return

    input_path = os.path.join(INPUT_FOLDER, files[0])
    print(f"📂 Loading file: {files[0]}")

    try:
        students, class_avg, internal_class = load_and_process_data(input_path)
    except Exception as e:
        print(f"❌ Error loading file: {e}")
        return

    if TEST_MODE:
        original_count = len(students)
        students = students[:MAX_STUDENTS]
        print(f"🧪 TEST MODE: Processing {len(students)} of {original_count} students")

    final_class_name = internal_class if internal_class else os.path.splitext(files[0])[0]
    print(f"📚 Class Name: {final_class_name}")
    print(f"👥 Total Students to Process: {len(students)}")
    print("-" * 50)

    for idx, student in enumerate(students, 1):
        name = student.get('name', 'Unknown')
        print(f"\n[{idx}/{len(students)}] Processing: {name}")

        try:
            analysis, meta = generate_llm_analysis(student, name, final_class_name)

            if analysis:
                create_word_doc(name, analysis, final_class_name, student, class_avg)
            else:
                print(f"❌ Failed to generate analysis for: {name}")

        except Exception as e:
            print(f"❌ Error processing {name}: {e}")
            continue

    print("\n" + "=" * 50)
    print("✅ Report generation complete!")
    print(f"📁 Output saved to: {OUTPUT_FOLDER}/")


if __name__ == "__main__":
    main()